# -*- coding: utf-8 -*-
"""Gera docs/ficha_visita_passo_a_passo.docx com a identidade Morabilidade.

O .docx abre direto no Google Docs (upload no Drive) preservando cores,
tabelas e callouts. Rodar com o Python do venv da API (tem python-docx):

    api\\.venv\\Scripts\\python.exe docs\\gerar_tutorial_docx.py
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RAIZ = Path(__file__).resolve().parent.parent
LOGO = RAIZ / "assets" / "logo.jpeg"
SAIDA = RAIZ / "docs" / "ficha_visita_passo_a_passo.docx"

OLIVE = "585a4f"
DOURADO = "d8cb6a"
SAND = "f7f6f2"
TEXTO = "333333"
TEXTO_SUAVE = "7a7c72"

FONTE = "Arial"


def _shade(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _borda_esquerda(par, hex_color: str, tamanho: int = 24) -> None:
    pPr = par._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(tamanho))
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _shade_par(par, hex_color: str) -> None:
    pPr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _run(par, texto: str, *, bold=False, cor=TEXTO, tamanho=10.5, italico=False):
    r = par.add_run(texto)
    r.font.name = FONTE
    r.font.size = Pt(tamanho)
    r.font.bold = bold
    r.font.italic = italico
    r.font.color.rgb = RGBColor.from_string(cor)
    return r


def secao(doc, titulo: str):
    par = doc.add_paragraph()
    par.space_before = Pt(14)
    par.paragraph_format.space_before = Pt(16)
    par.paragraph_format.space_after = Pt(2)
    _run(par, titulo.upper(), bold=True, cor=OLIVE, tamanho=13)
    # filete dourado sob o título
    pPr = par._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), DOURADO)
    pBdr.append(bottom)
    pPr.append(pBdr)


def corpo(doc, *trechos, espaco_depois=6):
    """trechos: lista de (texto, bold) ou string simples."""
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(espaco_depois)
    for t in trechos:
        if isinstance(t, tuple):
            _run(par, t[0], bold=t[1])
        else:
            _run(par, t)
    return par


def bullet(doc, *trechos):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.left_indent = Cm(0.6)
    par.paragraph_format.first_line_indent = Cm(-0.35)
    _run(par, "•  ", bold=True, cor=DOURADO, tamanho=10.5)
    for t in trechos:
        if isinstance(t, tuple):
            _run(par, t[0], bold=t[1])
        else:
            _run(par, t)
    return par


def numerado(doc, n: int, *trechos):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.left_indent = Cm(0.6)
    par.paragraph_format.first_line_indent = Cm(-0.35)
    _run(par, f"{n}.  ", bold=True, cor=OLIVE)
    for t in trechos:
        if isinstance(t, tuple):
            _run(par, t[0], bold=t[1])
        else:
            _run(par, t)
    return par


def callout(doc, texto: str):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after = Pt(8)
    par.paragraph_format.left_indent = Cm(0.3)
    _borda_esquerda(par, DOURADO)
    _shade_par(par, SAND)
    _run(par, texto, cor=OLIVE, tamanho=10, italico=True)
    return par


def main() -> None:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.6)
        s.bottom_margin = Cm(1.6)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)

    estilo = doc.styles["Normal"]
    estilo.font.name = FONTE
    estilo.font.size = Pt(10.5)

    # ── Faixa de marca (logo sobre fundo olive — mesmo tom do jpeg) ─────────
    faixa = doc.add_table(rows=1, cols=1)
    faixa.alignment = WD_TABLE_ALIGNMENT.CENTER
    cel = faixa.rows[0].cells[0]
    _shade(cel, OLIVE)
    p = cel.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.add_run().add_picture(str(LOGO), width=Cm(6.5))

    # ── Título ───────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(14)
    t.paragraph_format.space_after = Pt(2)
    _run(t, "FICHA DE VISITA", bold=True, cor=OLIVE, tamanho=20)
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_after = Pt(4)
    _run(st, "Passo a passo do corretor", cor=TEXTO_SUAVE, tamanho=12)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(8)
    intro.paragraph_format.space_after = Pt(4)
    _run(intro,
         "A ficha de visita é o termo digital que o visitante assina pelo celular antes de "
         "conhecer o imóvel. Ela protege a comissão (vincula o visitante à corretagem pelo "
         "prazo combinado, arts. 725/727 do Código Civil), cadastra o visitante "
         "automaticamente como cliente quando ele assina e monta o perfil de imóvel que ele busca.")

    # ── Antes de começar ─────────────────────────────────────────────────────
    secao(doc, "Antes de começar: acessar o painel")
    numerado(doc, 1, "Abra ", ("https://painel.morabilidade.com", True),
             " (funciona no computador e no navegador do celular).")
    numerado(doc, 2, "Entre com o ", ("e-mail e a senha", True), " do seu usuário.")
    numerado(doc, 3, "No menu lateral, clique em ", ("“Imóveis”", True),
             " e abra o imóvel que será visitado.")

    # ── 1. Gerar a ficha ─────────────────────────────────────────────────────
    secao(doc, "1. Gerar a ficha")
    numerado(doc, 1, "Abra o ", ("imóvel", True), " que será visitado no painel.")
    numerado(doc, 2, "Entre na aba ", ("“Fichas de visita”", True), ".")
    numerado(doc, 3, "Preencha os dados do visitante:")
    bullet(doc, ("Nome", True), " (obrigatório)")
    bullet(doc, ("WhatsApp", True),
           " (obrigatório — é por ele que o link é enviado e o cliente é cadastrado quando assina)")
    bullet(doc, ("CPF", True), " (recomendado — ajuda a não duplicar cadastros)")
    bullet(doc, ("E-mail", True), " (opcional)")
    numerado(doc, 4, "Clique em ", ("“Gerar e copiar link”", True), ".")
    corpo(doc, "O link de assinatura já fica copiado na área de transferência, "
               "pronto para colar em qualquer conversa.")

    # ── 2. Enviar ────────────────────────────────────────────────────────────
    secao(doc, "2. Enviar para o visitante")
    corpo(doc, "Na lista de ", ("fichas emitidas", True), ", use o botão de ",
          ("WhatsApp", True),
          " ao lado da ficha: ele abre a conversa com o visitante com a mensagem e o "
          "link prontos. Também dá para copiar o link de novo a qualquer momento.")
    callout(doc, "O link vale por 7 dias. Depois disso a ficha expira e é preciso gerar outra.")

    # ── 3. O visitante ───────────────────────────────────────────────────────
    secao(doc, "3. O que o visitante faz (no celular)")
    numerado(doc, 1, "Abre o link e confere os dados da visita (imóvel, endereço, valor, corretor).")
    numerado(doc, 2, "Lê a declaração e ", ("informa o CPF", True), ".")
    numerado(doc, 3, ("Assina com o dedo", True), " na tela e confirma o aceite.")
    corpo(doc, "Pronto: o sistema registra data/hora e IP, gera o ", ("PDF assinado", True),
          " e guarda tudo como prova. O visitante pode baixar o PDF na hora — e se abrir "
          "o link de novo depois, vê a confirmação com o botão de download.")

    # ── 4. Automático ────────────────────────────────────────────────────────
    secao(doc, "4. O que acontece sozinho (sem trabalho manual)")
    corpo(doc, "Tudo acontece ", ("quando o visitante assina", True),
          " (ficha que nunca é assinada não vira cadastro — a base de clientes só "
          "recebe quem realmente visitou):")
    bullet(doc, "O sistema procura o visitante na base de clientes pelo ",
           ("CPF, telefone ou e-mail", True), ".")
    bullet(doc, "Se já existe → a ficha é ", ("vinculada ao cadastro existente", True),
           " (nada é duplicado).")
    bullet(doc, "Se não existe → o visitante vira um ", ("cliente novo", True),
           " automaticamente, com origem “ficha de visita” e você como corretor responsável.")
    bullet(doc, "O CPF confirmado na assinatura completa o cadastro (se estava em branco).")
    bullet(doc, "O sistema monta o ", ("perfil de busca", True),
           " do cliente a partir de todas as visitas que ele já assinou: tipo de imóvel, "
           "cidade, bairros visitados, faixa de valor e dormitórios. Quanto mais visitas, "
           "mais preciso o perfil.")
    bullet(doc, "Esse perfil entra no módulo de ", ("Oportunidades", True),
           ": quando entrar um imóvel compatível, o cliente aparece como interessado.")

    # ── 5. Perfil ────────────────────────────────────────────────────────────
    secao(doc, "5. Perfil inferido × perfil manual")
    bullet(doc, "Na ficha do cliente, o perfil inferido aparece com um ", ("aviso âmbar", True),
           " indicando que foi calculado pelas visitas e que é recalculado a cada nova assinatura.")
    bullet(doc, "Se você ", ("editar e salvar", True), " o perfil, ele passa a ser ",
           ("manual", True), ": o sistema para de recalcular e respeita o que você definiu.")
    bullet(doc, "Um perfil que você já cadastrou manualmente ", ("nunca", True),
           " é sobrescrito pelas visitas.")

    # ── 6. Gerenciar ─────────────────────────────────────────────────────────
    secao(doc, "6. Gerenciar as fichas")
    corpo(doc, "Na aba “Fichas de visita” de cada imóvel:")

    tabela = doc.add_table(rows=4, cols=2)
    tabela.style = "Table Grid"
    larguras = (Cm(5.2), Cm(11.4))
    cabecalho = ("Ação", "Quando usar")
    linhas = [
        ("Copiar link / WhatsApp", "Reenviar o link enquanto a ficha está pendente."),
        ("Baixar PDF", "Pendente = rascunho para conferência; assinada = documento oficial "
                       "com trilha de auditoria."),
        ("Cancelar", "O link deixa de funcionar (não dá para cancelar ficha já assinada)."),
    ]
    for j, txt in enumerate(cabecalho):
        cel = tabela.rows[0].cells[j]
        cel.width = larguras[j]
        _shade(cel, OLIVE)
        p = cel.paragraphs[0]
        _run(p, txt, bold=True, cor="ffffff", tamanho=10)
    for i, (acao, quando) in enumerate(linhas, start=1):
        c0, c1 = tabela.rows[i].cells
        c0.width, c1.width = larguras
        if i % 2 == 0:
            _shade(c0, SAND)
            _shade(c1, SAND)
        _run(c0.paragraphs[0], acao, bold=True, cor=OLIVE, tamanho=10)
        _run(c1.paragraphs[0], quando, tamanho=10)

    fim = doc.add_paragraph()
    fim.paragraph_format.space_before = Pt(8)
    _run(fim, "Status possíveis: ", tamanho=10)
    _run(fim, "Aguardando assinatura", bold=True, cor=OLIVE, tamanho=10)
    _run(fim, " → ", tamanho=10)
    _run(fim, "Assinada", bold=True, cor=OLIVE, tamanho=10)
    _run(fim, " (ou ", tamanho=10)
    _run(fim, "Cancelada", bold=True, cor=TEXTO_SUAVE, tamanho=10)
    _run(fim, " / ", tamanho=10)
    _run(fim, "Expirada", bold=True, cor=TEXTO_SUAVE, tamanho=10)
    _run(fim, ").", tamanho=10)

    # ── 7. Boas práticas ─────────────────────────────────────────────────────
    secao(doc, "7. Boas práticas")
    bullet(doc, ("Sempre peça o CPF", True),
           " ao gerar a ficha: é o dado mais confiável para o sistema reconhecer um "
           "cliente que volta a visitar outros imóveis.")
    bullet(doc, "Gere a ficha ", ("antes", True),
           " da visita e envie pelo WhatsApp — o visitante assina no caminho ou na porta do imóvel.")
    bullet(doc, "Visitou de novo com outro corretor? Sem problema: o sistema reconhece o "
                "cliente pelo CPF/telefone e soma a nova visita ao mesmo perfil.")

    # ── Rodapé ───────────────────────────────────────────────────────────────
    rod = doc.sections[0].footer.paragraphs[0]
    rod.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(rod, "MORABILIDADE — Intermediação imobiliária  ·  painel.morabilidade.com",
         cor=TEXTO_SUAVE, tamanho=8)

    doc.save(SAIDA)
    print(f"OK: {SAIDA}")


if __name__ == "__main__":
    main()
